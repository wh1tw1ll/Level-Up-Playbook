import re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip separator lines
        if line.strip() in ('---', '___', '***'):
            i += 1
            continue
        
        # Headers
        if line.startswith('# ') or line.startswith('## '):
            level = 1 if line.startswith('# ') else 2
            text = re.sub(r'^#+\s*', '', line).strip()
            # Check for bold markers in header
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            h = doc.add_heading(text, level=level)
            if level == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Bold lines (like slide titles)
        if line.startswith('**') and line.endswith('**'):
            text = line.strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- '):
            text = line[2:]
            # Process inline bold
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatting(p, text)
            i += 1
            continue
        
        if line.startswith('  - '):
            text = re.sub(r'^\s+-\s*', '', line)
            # Indented sub-bullet
            p = doc.add_paragraph(style='List Bullet 2')
            _add_inline_formatting(p, text)
            i += 1
            continue
        
        # Numbered items (like "1. " or "One: ")
        numbered_match = re.match(r'^(One|Two|Three|Four|Five|Six|Seven|Eight):?\s+', line)
        if numbered_match:
            text = line
            p = doc.add_paragraph()
            _add_inline_formatting(p, text)
            i += 1
            continue
        
        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            _add_inline_formatting(p, line.strip())
        
        i += 1
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

def _add_inline_formatting(paragraph, text):
    """Add text with inline **bold** support."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.strip():
            run = paragraph.add_run(part)
            run.font.size = Pt(11)

if __name__ == '__main__':
    src = r'C:\Users\HermesAdmin\Level-Up-Playbook\pitches\chiefs\KC_Chiefs_Whitney_Script_v7_refined.md'
    dst = r'C:\Users\HermesAdmin\Level-Up-Playbook\pitches\chiefs\KC_Chiefs_Whitney_Script_v7_refined.docx'
    md_to_docx(src, dst)